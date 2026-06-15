from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 페이지 여백 설정 ──
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 기본 폰트 설정 ──
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def set_font(run, bold=False, size=10.5, color=None):
    run.font.name = '맑은 고딕'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text, color=(30, 80, 160)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=color)
    # 하단 보더
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'▶ {text}')
    set_font(run, bold=True, size=11.5, color=(46, 117, 182))
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True, size=10.5, color=(68, 68, 68))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        set_font(run, bold=True, size=10)
        # 헤더 배경색
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E75B6')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'EBF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].alignment = align
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=9.5)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg)
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ════════════════════════════════════════════
#  표지
# ════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('AI 바이브코딩 과정 교육 이수 보고서')
set_font(run, bold=True, size=22, color=(30, 80, 160))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AJW SCM 어시스턴트 개발 프로젝트')
set_font(run, bold=True, size=15, color=(68, 68, 68))

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    ('교육 과정', 'AI 바이브코딩 실무 활용 과정 (1~6회차)'),
    ('교육 기관', 'UPFLASH'),
    ('소속', '(주)에이제이월드 SCM팀'),
    ('성명', '박정원'),
    ('작성일', '2026년 6월 5일'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{label}:  ')
    set_font(r1, bold=True, size=11)
    r2 = p.add_run(value)
    set_font(r2, size=11)

doc.add_page_break()

# ════════════════════════════════════════════
#  1. 교육 개요
# ════════════════════════════════════════════
heading1('1. 교육 개요')

body('본 보고서는 AI 바이브코딩 실무 활용 과정(총 6회차) 이수 결과와, 해당 과정에서 습득한 기술을 실제 업무에 적용하여 개발한 「AJW SCM 어시스턴트」 프로젝트의 내용을 정리한 것입니다.')

doc.add_paragraph()
add_table(
    ['구분', '내용'],
    [
        ['교육 과정명', 'AI 바이브코딩 실무 활용 과정'],
        ['교육 기관',   'UPFLASH'],
        ['교육 기간',   '2026년 4월 ~ 2026년 6월 (총 6회차)'],
        ['교육 방식',   '강의 + 실습 (프로젝트 기반 학습)'],
        ['수강 목적',   'AI 코딩 도구를 활용한 SCM 업무 자동화 툴 개발'],
    ],
    col_widths=[3.5, 11]
)

# ════════════════════════════════════════════
#  2. 회차별 교육 내용
# ════════════════════════════════════════════
doc.add_paragraph()
heading1('2. 회차별 교육 내용')

sessions = [
    (
        '1회차 — AI와 바이브코딩 입문',
        [
            '바이브코딩(Vibe Coding) 개념 및 국내외 활용 사례 소개',
            '생성형 AI의 특성과 할루시네이션(Hallucination) 이해',
            '개발 환경 구성: Claude Desktop, Claude Code, VSCode, Git for Windows',
            '첫 번째 실습: 할일(Todo) 앱 만들기 — 프롬프트 구체성에 따른 결과 비교',
        ],
        '프로그래밍 지식 없이도 AI에게 명확하고 구체적인 지시를 내리는 것이 핵심임을 체득'
    ),
    (
        '2회차 — 환경 세팅과 MVP 개념',
        [
            '"손가락은 거들 뿐" — AI가 코드를 작성하고 사람은 방향을 제시하는 협업 방식 학습',
            'MVP(Minimum Viable Product) 개념: 핵심 기능만으로 빠르게 검증하는 개발 전략',
            '각자 실무 프로젝트 주제 선정 및 Claude Code 활용 기초 실습',
        ],
        '개인 프로젝트 주제를 「SCM 업무 자동화 웹 앱」으로 확정'
    ),
    (
        '3회차 — MVP 개발과 배포',
        [
            'Claude Code를 활용한 MVP 구현 실습',
            'GitHub를 통한 소스코드 버전 관리',
            'Vercel을 통한 웹 애플리케이션 배포 실습',
            '수강생 간 결과물 공유 및 피드백',
        ],
        'SCM 어시스턴트의 기본 탭 구조(생산자재 발주계획, 판매현황, 재고대사) MVP 완성'
    ),
    (
        '4회차 — 데이터베이스(DB) 연동',
        [
            'DB(Database) 개념 및 로컬 스토리지(localStorage) vs 서버 DB 비교',
            '스토리지 vs DB 용도 구분: 파일 저장과 데이터 관리의 차이',
            'Supabase를 활용한 클라우드 DB 연동 실습',
            'SQL 기초: 데이터 조회·삽입·수정·삭제',
        ],
        'Supabase를 프로젝트에 연동, 설정·품번·재고·판매 데이터의 클라우드 동기화 구현'
    ),
    (
        '5회차 — 외부 API 연동',
        [
            'API(Application Programming Interface) 개념: 서비스 간 데이터 연결 표준',
            'API 연결 구조: URL·인증키·요청(Request)·응답(Response) 흐름 이해',
            '한국수출입은행 현재환율 Open API 발급 및 프로젝트 연동 실습',
            '수강생 작품 시연: 해외 거래처 이슈 트래킹 시스템 등',
        ],
        '수익성 분석 기능에서 수입 원가 계산 시 환율 API 연동 필요성 확인'
    ),
    (
        '6회차 — RAG 적용과 프로젝트 아키텍처 정비',
        [
            'RAG(Retrieval-Augmented Generation) 개념: AI가 내부 데이터를 실시간 참조하여 답변하는 기술',
            'LLM → RAG → Advanced RAG → Graph RAG → Agentic RAG 기술 발전 단계 이해',
            'Claude Code MD 파일 체계화: CLAUDE.md / ARCHITECTURE.md / README.md 역할 구분',
            'CLAUDE.md 계층 구조 활용법: 글로벌(~/.claude) / 프로젝트 / 앱 폴더 단위 분리',
            'Git 버전 관리 심화: git log / git checkout / git revert / git reset 활용',
            'Auto Mode와 Context Window 관리를 통한 효율적인 AI 협업 방법론',
        ],
        'ARCHITECTURE.md 작성 완료, RAG 적용 계획 수립'
    ),
]

for title, points, outcome in sessions:
    heading2(title)
    for pt in points:
        bullet(pt)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run('→ 프로젝트 적용: ')
    set_font(r1, bold=True, size=10, color=(46, 117, 182))
    r2 = p.add_run(outcome)
    set_font(r2, size=10)

# ════════════════════════════════════════════
#  3. 개발 프로젝트 소개
# ════════════════════════════════════════════
doc.add_page_break()
heading1('3. 개발 프로젝트: AJW SCM 어시스턴트')

body('(주)에이제이월드 SCM팀의 업무를 지원하는 웹 기반 공급망 관리(SCM) 도구입니다. 완제품 수입 관리, 생산자재 발주계획, 판매현황 분석, 재고 대사 등 기존에 수작업으로 처리하던 Excel 업무를 자동화·시각화합니다.')

doc.add_paragraph()
heading3('기술 스택')
add_table(
    ['분류', '기술'],
    [
        ['프론트엔드',  'React 19, TypeScript, Vite, Tailwind CSS 4'],
        ['Excel 처리', 'ExcelJS, SheetJS (xlxs)'],
        ['데이터 저장', 'localStorage, IndexedDB, Supabase (PostgreSQL)'],
        ['개발 도구',   'Claude Code, VSCode, Git'],
        ['배포 환경',   'localhost:5173 (개발 서버)'],
    ],
    col_widths=[3.5, 11]
)

doc.add_paragraph()
heading3('주요 기능 구성 (7개 탭)')
add_table(
    ['탭', '기능 설명', '상태'],
    [
        ['홈 대시보드',     'KPI 6개·수입 D-day·재고대사 현황·OJC CAGR 시각화',           '완료'],
        ['생산자재 발주계획', 'ERP 파싱(STEP1) → 판매분석(STEP2) → 발주계획 생성(STEP3)',   '개발중'],
        ['판매현황 분석',   '73,000건+ 판매 데이터 파싱·연도별 분석·8시트 Excel 다운로드',   '완료'],
        ['재고 대사',       'EMP ↔ 이카운트 재고 비교·이력 누적·추이 차트',                 '완료'],
        ['수입 관리',       '발주계획·간헐적수요·발주현황(D-day)·수익성 분석',               '완료'],
        ['품번 생성기',     'OJC 품번 자동/수동 생성, EMP 로트코드 생성',                   '개발중'],
        ['관리자',          '자재관리·OJC 코드표 편집·설정 (비밀번호 게이트)',               '완료'],
    ],
    col_widths=[3.2, 9.5, 1.8]
)

doc.add_paragraph()
heading3('핵심 비즈니스 로직')
bullet('안전재고 산출: 월평균 수요 × (리드타임 / 30일) × 안전계수(k)')
bullet('발주필요량: max(0, 예측연간수요 + 안전재고 − 현재고 − 기발주)')
bullet('수익성 분석 마진율: (판매가 − 원가) / 판매가 × 100')
bullet('맥산 생산 가중치: FLC 대비 최대 5% 비싸도 국내 생산 유지 (CEO 지시사항 반영)')
bullet('OJC 품번 자동 분류: 정규화된 품목명 → KT/LG OJC·DROP·피그테일 등 카테고리 매핑')

# ════════════════════════════════════════════
#  4. AI 활용 방법론
# ════════════════════════════════════════════
doc.add_paragraph()
heading1('4. AI 활용 방법론')

heading2('Claude Code 중심 개발 방식')
bullet('전체 코드의 100%를 Claude Code와의 대화를 통해 작성')
bullet('CLAUDE.md에 도메인 지식·코딩 컨벤션·타입 정의를 문서화하여 AI 컨텍스트 품질 향상')
bullet('ARCHITECTURE.md를 통해 컴포넌트 계층·데이터 흐름·저장소 구조를 AI와 공유')
bullet('Auto Mode 활용으로 반복적인 수정 작업 자동화')

heading2('바이브코딩 철학의 실천')
bullet('1회차 핵심 교훈 "구체성의 차이"를 실무에 적용: 업무 도메인(OJC 품번 체계, 이카운트-EMP 병행 구조)을 프롬프트에 명시')
bullet('AI가 생성한 코드를 검토·수정하는 역할 분담: AI가 구현, 담당자가 비즈니스 로직 검증')
bullet('기능 단위 MVP 방식으로 탭별 순차 개발, 피드백 반영 후 다음 기능으로 진행')

heading2('데이터 관리 전략 (4회차 학습 적용)')
bullet('localStorage: 즉시 반응성이 필요한 모든 상태의 1차 캐시')
bullet('Supabase: 팀 공유가 필요한 마스터 데이터(품번·재고·설정) 동기화')
bullet('IndexedDB: 73,000건 이상의 판매 상세 데이터 대용량 저장')

# ════════════════════════════════════════════
#  5. 업무 적용 효과
# ════════════════════════════════════════════
doc.add_paragraph()
heading1('5. 업무 적용 효과')

add_table(
    ['업무 영역', '기존 방식', '개선 후', '효과'],
    [
        ['완제품 발주계획',  '수작업 Excel 집계',      '파일 업로드 → 자동 계산 → Excel 출력', '작업 시간 대폭 단축'],
        ['수익성 분석',     '원가/판매가 수동 비교',   '원가 파일 업로드 → 마진율 자동 산출',   '품목별 즉시 판단 가능'],
        ['재고 대사',       'EMP·이카운트 수동 대조', '파일 4종 업로드 → 차이 자동 계산',      '이력 누적·추이 분석'],
        ['판매 데이터 분석', '다수 시트 수동 정리',    '73,000건 자동 파싱·카테고리별 집계',    '연도별 CAGR 자동 산출'],
        ['품번 관리',       '담당자 개인 관리',        '웹 기반 코드표·품번 생성기',            '팀 공유 및 일관성 확보'],
    ],
    col_widths=[2.8, 3.5, 4.5, 3.7]
)

# ════════════════════════════════════════════
#  6. 향후 계획
# ════════════════════════════════════════════
doc.add_paragraph()
heading1('6. 향후 개발 계획')

heading2('단기 (1~2개월)')
bullet('RAG 기능 구현: 판매·재고 데이터를 Claude API와 연동하여 자연어 질의응답 기능 추가')
bullet('  - "이번 달 발주 급한 품목은?", "C-ALNN-3 재고 며칠치 남았어?" 등 질의 가능', level=1)
bullet('한국수출입은행 환율 API 연동: 수익성 분석 시 실시간 환율 자동 반영')
bullet('Vercel 배포: 팀 전체가 접근 가능한 공개 URL 확보')

heading2('중장기 (3~6개월)')
bullet('EMP 로케이션 KT향 재구성: 창고 위치 자동 배정 로직 구현')
bullet('리드타임 집계 분석: 발주~입고 실적 데이터 기반 리드타임 통계')
bullet('모바일 반응형 UI 개선: 현장에서 스마트폰으로 재고 조회 가능')

# ════════════════════════════════════════════
#  7. 총평 및 소감
# ════════════════════════════════════════════
doc.add_paragraph()
heading1('7. 총평 및 소감')

body('6회차에 걸친 AI 바이브코딩 교육을 통해, 프로그래밍 전문 지식 없이도 실무에서 즉시 활용 가능한 수준의 웹 애플리케이션을 완성할 수 있었습니다.')
doc.add_paragraph()
body('특히 다음 세 가지가 이번 교육의 핵심 성과였습니다.')
bullet('구체성의 힘: 업무 도메인 지식을 프롬프트에 담을수록 AI의 결과물 품질이 높아진다는 것을 직접 체험')
bullet('문서의 중요성: CLAUDE.md·ARCHITECTURE.md 등 구조화된 문서가 AI와의 협업 품질을 결정적으로 좌우함을 확인')
bullet('빠른 검증: MVP 방식으로 탭 단위로 빠르게 만들고 실무 피드백을 받아 개선하는 사이클이 효과적임을 체감')
doc.add_paragraph()
body('현재 동료들이 실제 업무에서 SCM 어시스턴트를 사용하고 있으며, 향후 RAG 기능과 환율 API를 추가하여 더욱 완성도 높은 업무 자동화 도구로 발전시킬 계획입니다.')

# ── 저장 ──
out_path = r'C:\Users\AJWorld\todo-app\AJW_바이브코딩_교육보고서_박정원.docx'
doc.save(out_path)
print(f'저장 완료: {out_path}')
